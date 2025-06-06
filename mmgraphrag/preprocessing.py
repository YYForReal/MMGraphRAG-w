from typing import Callable, Dict, List, Optional, Type, Union, cast
from dataclasses import dataclass
from docx import Document
from PIL import Image
from io import BytesIO
from parameter import cache_path
import asyncio
import os
import json
import base64

from base import (
    logger,
    compute_mdhash_id,
    encode_string_by_tiktoken,
    decode_tokens_by_tiktoken,
    read_config_to_dict,
)

from storage import (
    BaseKVStorage,
    JsonKVStorage,
    StorageNameSpace,
)
from llm import multimodel_if_cache
from prompt import PROMPTS

global_config_path = os.path.join(cache_path, "global_config.csv")

def chunking_by_token_size(
    content: str, overlap_token_size=128, max_token_size=1024, tiktoken_model="gpt-4o"
):
    """
    根据token大小对文本进行分块。

    该函数用于将给定的文本内容按照指定的token大小限制进行分块，同时保证相邻块之间有重叠。
    主要用于处理大文本，使其能够适应如OpenAI的GPT系列模型的输入限制。

    参数:
    - content: str, 待分块的文本内容。
    - overlap_token_size: int, 默认128. 相邻文本块之间的重叠token数。
    - max_token_size: int, 默认1024. 每个文本块的最大token数。
    - tiktoken_model: str, 默认"gpt-4o". 用于token化和去token化的tiktoken模型。

    返回:
    - List[Dict[str, Any]], 包含每个文本块的tokens数量、文本内容和块顺序索引的列表。
    """
    # 使用指定的tiktoken模型对文本进行token化
    tokens = encode_string_by_tiktoken(content, model_name=tiktoken_model)
    # 初始化存储分块结果的列表
    results = []
    # 遍历tokens，根据max_token_size和overlap_token_size进行分块
    for index, start in enumerate(
        range(0, len(tokens), max_token_size - overlap_token_size)
    ):
        # 根据当前分块的起始位置和最大token数限制，获取分块的tokens
        chunk_content = decode_tokens_by_tiktoken(
            tokens[start : start + max_token_size], model_name=tiktoken_model
        )
        # 将当前分块的tokens数量、文本内容和块顺序索引添加到结果列表中
        results.append(
            {
                "tokens": min(max_token_size, len(tokens) - start),
                "content": chunk_content.strip(),
                "chunk_order_index": index,
            }
        )
    return results

@dataclass
class text_chunking_func:
    # 分块
    chunk_func: Callable[[str, Optional[int], Optional[int], Optional[str]], List[Dict[str, Union[str, int]]]] = chunking_by_token_size
    # 分块大小
    chunk_token_size: int = 1200
    # 分块重叠数量
    chunk_overlap_token_size: int = 100
    

    # 键值存储，json，具体定义在storage.py
    key_string_value_json_storage_cls: Type[BaseKVStorage] = JsonKVStorage
    # 获取全局设置
    global_config = read_config_to_dict(global_config_path)
    
    # tiktoken使用的模型名字，默认为gpt-4o，moonshot-v1-32k可以通用
    tiktoken_model_name = global_config["tiktoken_model_name"]

    def __post_init__(self):
        # 初始化存储类实例，用于存储完整文档
        self.full_docs = self.key_string_value_json_storage_cls(
            namespace="full_docs", global_config = self.global_config
        )
        # 初始化存储类实例，用于存储文本块
        self.text_chunks = self.key_string_value_json_storage_cls(
            namespace="text_chunks", global_config = self.global_config
        )
    
    async def text_chunking(self,string_or_strings):
        try:
            # 如果输入是一个字符串，将其转换为列表
            if isinstance(string_or_strings, str):
                string_or_strings = [string_or_strings]
            # ---------- new docs
            # 将字符串或字符串列表string_or_strings中的每个元素去除首尾空白后，作为文档内容。
            # 计算其MD5哈希值并添加前缀doc-作为键，内容本身作为值，生成一个新的字典new_docs。
            new_docs = {
                compute_mdhash_id(c.strip(), prefix="doc-"): {"content": c.strip()}
                for c in string_or_strings
            }
            # 筛选出需要添加的新文档ID
            _add_doc_keys = await self.full_docs.filter_keys(list(new_docs.keys()))
            # 根据筛选结果更新新文档字典
            new_docs = {k: v for k, v in new_docs.items() if k in _add_doc_keys}
            # 如果没有新文档需要添加，记录日志并返回
            if not len(new_docs):
                logger.warning(f"All docs are already in the storage")
                return
            # 记录插入新文档的日志
            logger.info(f"[New Docs] inserting {len(new_docs)} docs")

            # ---------- chunking
            inserting_chunks = {}
            for doc_key, doc in new_docs.items():
                # 为每个文档生成片段
                chunks = {
                    compute_mdhash_id(dp["content"], prefix="chunk-"): {
                        **dp,
                        "full_doc_id": doc_key,
                    }
                    for dp in self.chunk_func(
                        doc["content"],
                        overlap_token_size=self.chunk_overlap_token_size,
                        max_token_size=self.chunk_token_size,
                        tiktoken_model=self.tiktoken_model_name,
                    )
                }
                inserting_chunks.update(chunks)
            # 筛选出需要添加的新片段ID
            _add_chunk_keys = await self.text_chunks.filter_keys(
                list(inserting_chunks.keys())
            )
            # 根据筛选结果更新新片段字典
            inserting_chunks = {
                k: v for k, v in inserting_chunks.items() if k in _add_chunk_keys
            }
            # 如果没有新片段需要添加，记录日志并返回
            if not len(inserting_chunks):
                logger.warning(f"All chunks are already in the storage")
                return
            logger.info(f"[New Chunks] inserting {len(inserting_chunks)} chunks")

            # 提交所有更新和索引操作
            await self.full_docs.upsert(new_docs)
            await self.text_chunks.upsert(inserting_chunks)
        finally:
            await self._text_chunking_done()
    async def _text_chunking_done(self):
        tasks = []
        for storage_inst in [
            self.full_docs,
            self.text_chunks,
        ]:
            if storage_inst is None:
                continue
            tasks.append(cast(StorageNameSpace, storage_inst).index_done_callback())
        await asyncio.gather(*tasks)

image_description_prompt_user = PROMPTS["image_description_user"]
image_description_prompt_system = PROMPTS["image_description_system"]

async def get_image_description(image_path):
    with open(image_path, 'rb') as img_file:
        image_base = base64.b64encode(img_file.read()).decode('utf-8')
    image_description =  await multimodel_if_cache(user_prompt=image_description_prompt_user,img_base=image_base,system_prompt=image_description_prompt_system)
    return image_description

def find_chunk_for_image(text_chunks, preceding_text, following_text):
    """
    根据图片前后文本找到其所属 chunk。
    优先选择包含更多连续字符的 chunk，忽略换行符。
    """
    best_chunk_id = None
    best_match_count = 0

    # 将前后文本合并为一个整体，去掉换行符
    combined_text = f"{preceding_text} {following_text}".replace('\n', '').strip()

    # 如果合并的文本为空，则返回 None
    if not combined_text:
        return None

    # 遍历所有的 chunk
    for chunk_id, chunk_data in text_chunks.items():
        # 去掉 chunk 中的换行符
        chunk_content = chunk_data['content'].replace('\n', '')

        # 计算组合文本与 chunk 内容的匹配度（基于词语匹配）
        match_count = sum(1 for word in combined_text.split() if word in chunk_content)

        # 如果当前 chunk 的匹配度最高，则选择它
        if match_count > best_match_count:
            best_match_count = match_count
            best_chunk_id = chunk_id

    return best_chunk_id

def extract_image_context(docx_path, context_length=100):
    """
    提取文档中的图片前后文本信息。
    """
    doc = Document(docx_path)
    results = {}
    image_counter = 1  # 初始化图像计数器

    def get_previous_text(index, length_needed):
        """
        从前面的段落中递归获取文本，直到满足length_needed字符数。
        """
        text = ""
        while index >= 0 and len(text) < length_needed:
            paragraph_text = doc.paragraphs[index].text
            text = paragraph_text[-(length_needed - len(text)):] + text
            index -= 1
        return text

    def get_next_text(index, length_needed):
        """
        从后面的段落中递归获取文本，直到满足length_needed字符数。
        """
        text = ""
        while index < len(doc.paragraphs) and len(text) < length_needed:
            paragraph_text = doc.paragraphs[index].text
            text = text + paragraph_text[:length_needed - len(text)]
            index += 1
        return text

    for i, paragraph in enumerate(doc.paragraphs):
        # 获取每个段落的文本
        text = paragraph.text

        # 检查该段落是否包含图像
        has_image = any(run._element.xpath('.//a:blip') for run in paragraph.runs)

        if has_image:
            # 提取图像前后文本

            # 获取前面的段落文本
            before_text = get_previous_text(i - 1, context_length)

            # 获取后面的段落文本
            after_text = get_next_text(i + 1, context_length)

            # 保存结果，确保前后文本不为空
            results[f"image_{image_counter}"] = {
                "before": before_text.strip(),
                "after": after_text.strip()
            }

            # 递增图像计数器
            image_counter += 1

    return results

def compress_image_to_size(input_image, output_image_path, target_size_mb=5, step=10, quality=90):
    """
    将图片压缩到目标大小以内（以MB为单位）。

    参数:
    input_image (PIL.Image): 输入图片的 PIL 对象。
    output_image_path (str): 输出图片的路径。
    target_size_mb (int): 目标大小，以MB为单位，默认为5MB。
    step (int): 每次降低的质量步长，默认为10。
    quality (int): 初始保存的图片质量，默认为90。

    返回:
    bool: 是否成功压缩到目标大小以内。
    """
    target_size_bytes = target_size_mb * 1024 * 1024  # 将目标大小转换为字节

    # 先保存图片并检查大小
    img = input_image
    img.save(output_image_path, quality=quality)
    if os.path.getsize(output_image_path) <= target_size_bytes:
        return True

    # 尝试逐步降低质量，直到图片大小小于目标大小
    while os.path.getsize(output_image_path) > target_size_bytes and quality > 10:
        quality -= step
        img.save(output_image_path, quality=quality)
    
    # 检查最终大小是否符合要求
    if os.path.getsize(output_image_path) <= target_size_bytes:
        return True
    else:
        print("无法将图片压缩到目标大小以内，请在preprocessing.py中调整初始质量或步长。")
        return False

async def extract_text_and_images_with_chunks(docx_path, output_dir,context_length):
    """
    提取文档中的文本块，并与图片关联。图片前后上下文文本提取整合。
    """
    doc = Document(docx_path)

    # 提取文档的所有文本内容
    full_text = "\n".join([para.text for para in doc.paragraphs])

    # 首先实例化类
    text_chunking_instance = text_chunking_func()
    # 对文档文本进行分块
    await text_chunking_instance.text_chunking(full_text)
    filepath = os.path.join(output_dir, 'kv_store_text_chunks.json')
    with open(filepath, 'r') as file:
        text_chunks = json.load(file)

    # 创建保存图片的目录
    images_dir = os.path.join(output_dir, 'images')
    if not os.path.exists(images_dir):
        os.makedirs(images_dir)

    # 提取图像上下文（前后文本）
    image_contexts = extract_image_context(docx_path, context_length)

    # 生成最终的结果字典
    image_data = {}
    image_count = 0
    extracted_images = []  # 追踪提取到的图片

    # Step 1: 使用 XML 解析提取图片（按顺序）
    for idx, shape in enumerate(doc.element.xpath("//w:drawing//a:blip"), start=1):
        # 提取 `r:embed` 属性，它指向图片的关系 ID
        embed = shape.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if embed in doc.part.rels:
            image_count += 1
            image = doc.part.rels[embed].target_part.blob

            # 使用 PIL 打开图片，并将其转换为 JPG 格式
            image_bytes = BytesIO(image)
            with Image.open(image_bytes) as img:
                rgb_img = img.convert('RGB')  # 转换为 RGB 模式，适用于 JPG 格式
                image_output_path = os.path.join(images_dir, f'image_{image_count}.jpg')
                # 使用压缩功能，将图片保存为小于 target_size_mb 的大小
                compress_image_to_size(rgb_img, image_output_path, target_size_mb=5)

            # 记录图片信息，用于后续处理
            extracted_images.append({
                'image_id': image_count,
                'image_path': image_output_path,
                'image_description': await get_image_description(image_output_path)
            })

    # Step 2: 关联图片上下文
    for image_key, context in image_contexts.items():
        # 从 image_key 提取图片索引（避免超出索引）
        image_index = int(image_key.split('_')[1]) - 1
        if image_index < len(extracted_images):
            current_image = extracted_images[image_index]  # 获取当前图片
            preceding_text_segment = context["before"]
            following_text_segment = context["after"]
            
            
            # 找到图片所属的 chunk
            associated_chunk_id = find_chunk_for_image(text_chunks, preceding_text_segment, following_text_segment)
            if associated_chunk_id:
                # 将图片信息添加到 image_data 中
                image_data[image_key] = {
                    "chunk_order_index": text_chunks[associated_chunk_id]['chunk_order_index'],
                    "chunk_id": associated_chunk_id,
                    "image_id": current_image['image_id'],
                    "description": current_image['image_description'],
                    "image_path": current_image['image_path'],
                    "before": preceding_text_segment,
                    "after": following_text_segment
                }
    return image_data

@dataclass
class chunking_func:
    # 图像提取上下文长度（各100，所以总长度为200）
    context_length: int = 100

    # 键值存储，json，具体定义在storage.py
    key_string_value_json_storage_cls: Type[BaseKVStorage] = JsonKVStorage

    # 获取全局设置
    global_config = read_config_to_dict(global_config_path)

    def __post_init__(self):
        # 初始化存储类实例，用于存储图像属性
        self.image_data = self.key_string_value_json_storage_cls(
            namespace="image_data", global_config = self.global_config
        )
    
    async def extract_text_and_images(self,docx_path):
        try:
            output_dir = self.global_config["working_dir"]
            context_length = self.context_length
            imagedata = await extract_text_and_images_with_chunks(docx_path, output_dir, context_length)
            # 提交所有更新和索引操作
            await self.image_data.upsert(imagedata)
        finally:
            await self._chunking_done()
    async def _chunking_done(self):
        tasks = []
        for storage_inst in [
            self.image_data
        ]:
            if storage_inst is None:
                continue
            tasks.append(cast(StorageNameSpace, storage_inst).index_done_callback())
        await asyncio.gather(*tasks)